"""
File Upload Component
Handles resume file uploads with validation and error handling
"""

import streamlit as st
import io
from typing import Optional

# Unicode symbols for UI
SYMBOLS = {
    'upload': '📄',
    'success': '✅',
    'error': '❌',
    'warning': '⚠️',
    'info': 'ℹ️',
    'document': '📋',
    'check': '✓',
    'cross': '✗'
}

def handle_file_upload() -> Optional[st.runtime.uploaded_file_manager.UploadedFile]:
    """Handle file upload with validation"""

    # File upload widget
    uploaded_file = st.file_uploader(
        f"{SYMBOLS['upload']} Choose your resume file",
        type=['pdf', 'docx', 'txt'],
        help="Supported formats: PDF, DOCX, TXT (Max size: 10MB)",
        key="resume_uploader"
    )

    if uploaded_file is not None:
        # Validate file
        validation_result = validate_uploaded_file(uploaded_file)

        if validation_result['valid']:
            # Store in session state immediately
            st.session_state['uploaded_file'] = uploaded_file
            st.session_state['file_name'] = uploaded_file.name

            # Reset processing flags for new files
            if st.session_state.get('last_processed_file') != uploaded_file.name:
                st.session_state['content_extracted'] = False
                st.session_state['last_processed_file'] = uploaded_file.name
                st.session_state['processing_stage'] = 'uploaded'

            # Show success message
            st.success(f"{SYMBOLS['success']} File uploaded successfully!")

            # Show file details
            with st.expander(f"{SYMBOLS['info']} File Details", expanded=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**Name:** {uploaded_file.name}")
                    st.write(f"**Size:** {format_file_size(uploaded_file.size)}")
                with col2:
                    st.write(f"**Type:** {uploaded_file.type}")
                    st.write(f"**Status:** {SYMBOLS['check']} Ready to process")

            return uploaded_file
        else:
            # Show validation errors
            st.error(f"{SYMBOLS['error']} File validation failed:")
            for error in validation_result['errors']:
                st.error(f"  • {error}")
            return None

    # Show upload instructions when no file is selected
    else:
        st.info(f"""
        {SYMBOLS['info']} **Upload Instructions:**

        • **Supported formats:** PDF, DOCX, TXT
        • **Maximum size:** 10MB
        • **Best quality:** DOCX files typically provide the best text extraction
        • **Tips:** Ensure your resume is not password-protected or corrupted
        """)

    return None

def validate_uploaded_file(uploaded_file) -> dict:
    """Validate uploaded file"""
    errors = []

    # Check file size (10MB limit)
    max_size = 10 * 1024 * 1024  # 10MB in bytes
    if uploaded_file.size > max_size:
        errors.append(f"File size ({format_file_size(uploaded_file.size)}) exceeds 10MB limit")

    # Check if file is empty
    if uploaded_file.size == 0:
        errors.append("File appears to be empty")

    # Check file type
    allowed_types = [
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'text/plain'
    ]

    if uploaded_file.type not in allowed_types:
        errors.append(f"Unsupported file type: {uploaded_file.type}")

    # Additional checks for specific file types
    if uploaded_file.type == 'application/pdf':
        errors.extend(validate_pdf_file(uploaded_file))
    elif 'wordprocessingml' in uploaded_file.type:
        errors.extend(validate_docx_file(uploaded_file))

    return {
        'valid': len(errors) == 0,
        'errors': errors
    }

def validate_pdf_file(uploaded_file) -> list:
    """Validate PDF file"""
    errors = []

    try:
        # Reset file pointer
        uploaded_file.seek(0)

        # Try to read PDF
        import PyPDF2
        pdf_reader = PyPDF2.PdfReader(uploaded_file)

        # Check if PDF has pages
        if len(pdf_reader.pages) == 0:
            errors.append("PDF file contains no pages")

        # Check if PDF is encrypted
        if pdf_reader.is_encrypted:
            errors.append("PDF file is password-protected. Please upload an unprotected version.")

        # Reset file pointer after validation
        uploaded_file.seek(0)

    except Exception as e:
        errors.append(f"PDF file appears to be corrupted: {str(e)}")

    return errors

def validate_docx_file(uploaded_file) -> list:
    """Validate DOCX file"""
    errors = []

    try:
        # Reset file pointer
        uploaded_file.seek(0)

        # Try to read DOCX
        import docx
        doc = docx.Document(uploaded_file)

        # Check if document has content
        has_content = False
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                has_content = True
                break

        if not has_content:
            # Check tables as well
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            has_content = True
                            break
                    if has_content:
                        break
                if has_content:
                    break

        if not has_content:
            errors.append("DOCX file appears to be empty or contains no readable text")

        # Reset file pointer after validation
        uploaded_file.seek(0)

    except Exception as e:
        errors.append(f"DOCX file appears to be corrupted: {str(e)}")

    return errors

def format_file_size(size_bytes: int) -> str:
    """Format file size in human readable format"""
    if size_bytes == 0:
        return "0 B"

    size_names = ["B", "KB", "MB", "GB"]
    import math
    i = int(math.floor(math.log(size_bytes, 1024)))
    p = math.pow(1024, i)
    s = round(size_bytes / p, 2)
    return f"{s} {size_names[i]}"

def clear_uploaded_file():
    """Clear uploaded file from session state"""
    keys_to_clear = [
        'uploaded_file', 'file_name', 'content_extracted',
        'last_processed_file', 'extracted_content', 'enhanced_resume',
        'cover_letter', 'processing_stage'
    ]

    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]

    st.success(f"{SYMBOLS['success']} File cleared successfully!")
    st.experimental_rerun()

def get_file_preview(uploaded_file, max_chars: int = 500) -> str:
    """Get a preview of the file content"""
    try:
        if uploaded_file.type == 'text/plain':
            uploaded_file.seek(0)
            content = uploaded_file.read().decode('utf-8', errors='ignore')
            uploaded_file.seek(0)
            return content[:max_chars] + ("..." if len(content) > max_chars else "")
        else:
            return f"Preview not available for {uploaded_file.type} files. File will be processed when you click 'Process Resume'."
    except Exception:
        return "Preview not available."
