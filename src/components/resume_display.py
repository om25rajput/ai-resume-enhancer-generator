"""
Resume Display Component
Shows original and enhanced resume content
"""

import streamlit as st
from typing import Dict, Any
import io

# Unicode symbols for UI
SYMBOLS = {
    'document': '📋',
    'magic': '✨',
    'download': '⬇️',
    'star': '⭐',
    'check': '✅',
    'info': 'ℹ️',
    'warning': '⚠️',
    'arrow_right': '→',
    'bullet': '•',
    'gear': '⚙️'
}

def display_resume(uploaded_file):
    """Display resume content and enhancements"""

    # Tabs for different views
    tab1, tab2, tab3, tab4 = st.tabs([
        f"{SYMBOLS['document']} Original", 
        f"{SYMBOLS['magic']} Enhanced", 
        f"{SYMBOLS['star']} Analysis",
        f"{SYMBOLS['download']} Export"
    ])

    with tab1:
        display_original_content(uploaded_file)

    with tab2:
        display_enhanced_content()

    with tab3:
        display_analysis()

    with tab4:
        display_export_options()

def display_original_content(uploaded_file):
    """Display original resume content"""
    st.markdown(f"### {SYMBOLS['document']} Original Resume")

    if 'extracted_content' in st.session_state and st.session_state.extracted_content:
        content = st.session_state.extracted_content

        # Show content preview
        st.markdown("#### Content Preview")
        st.text_area(
            "Extracted Text",
            content[:2000] + ("..." if len(content) > 2000 else ""),
            height=300,
            disabled=True
        )

        # Show statistics
        col1, col2, col3 = st.columns(3)

        with col1:
            st.metric("Characters", len(content))

        with col2:
            word_count = len(content.split())
            st.metric("Words", word_count)

        with col3:
            line_count = len(content.split('\n'))
            st.metric("Lines", line_count)

    else:
        st.info(f"{SYMBOLS['info']} Upload and process a resume to see the original content here.")

def display_enhanced_content():
    """Display AI-enhanced resume content"""
    st.markdown(f"### {SYMBOLS['magic']} AI-Enhanced Resume")

    if 'enhanced_resume' in st.session_state and st.session_state.enhanced_resume:
        enhanced = st.session_state.enhanced_resume

        # Enhanced summary
        if enhanced.get('enhanced_summary'):
            st.markdown("#### Professional Summary")
            st.markdown(f'<div style="background-color: #f0f2f6; padding: 1rem; border-radius: 0.5rem; border-left: 4px solid #007bff;">{enhanced["enhanced_summary"]}</div>', unsafe_allow_html=True)

        # Enhanced skills
        if enhanced.get('enhanced_skills'):
            st.markdown("#### Enhanced Skills")
            skills = enhanced['enhanced_skills']

            if skills.get('technical'):
                st.markdown("**Technical Skills:**")
                st.write(", ".join(skills['technical']))

            if skills.get('soft_skills'):
                st.markdown("**Soft Skills:**")
                st.write(", ".join(skills['soft_skills']))

            if skills.get('suggested_additions'):
                st.markdown("**Suggested Additions:**")
                for skill in skills['suggested_additions']:
                    st.markdown(f"{SYMBOLS['star']} {skill}")

        # Enhanced experience
        if enhanced.get('enhanced_experience'):
            st.markdown("#### Enhanced Experience")
            for exp in enhanced['enhanced_experience']:
                st.markdown(f"**{exp.get('role', 'Role')}** - {exp.get('company', 'Company')}")
                for desc in exp.get('enhanced_description', []):
                    st.markdown(f"{SYMBOLS['bullet']} {desc}")
                st.markdown("")

        # Full enhanced content
        if enhanced.get('enhanced_full_content'):
            st.markdown("#### Complete Enhanced Resume")
            st.text_area(
                "Enhanced Resume Content",
                enhanced['enhanced_full_content'],
                height=400,
                disabled=True
            )

    else:
        st.info(f"{SYMBOLS['info']} Process your resume to see AI enhancements here.")

def display_analysis():
    """Display resume analysis and suggestions"""
    st.markdown(f"### {SYMBOLS['star']} Resume Analysis")

    if 'enhanced_resume' in st.session_state and st.session_state.enhanced_resume:
        enhanced = st.session_state.enhanced_resume

        # Improvement suggestions
        if enhanced.get('suggested_improvements'):
            st.markdown("#### Improvement Suggestions")
            for suggestion in enhanced['suggested_improvements']:
                if suggestion.strip():
                    st.markdown(suggestion)

        # ATS optimization
        if enhanced.get('ats_optimizations'):
            st.markdown("#### ATS Optimization")
            ats = enhanced['ats_optimizations']

            # ATS Score
            if ats.get('overall_ats_score'):
                score = ats['overall_ats_score']
                color = "green" if score >= 80 else "orange" if score >= 60 else "red"
                st.markdown(f"**ATS Compatibility Score:** <span style='color: {color}; font-weight: bold;'>{score}/100</span>", unsafe_allow_html=True)

            # Keywords to add
            if ats.get('keywords_to_add'):
                st.markdown("**Keywords to Add:**")
                for keyword in ats['keywords_to_add']:
                    st.markdown(f"{SYMBOLS['star']} {keyword}")

            # Formatting improvements
            if ats.get('formatting_improvements'):
                st.markdown("**Formatting Improvements:**")
                for improvement in ats['formatting_improvements']:
                    st.markdown(f"{SYMBOLS['check']} {improvement}")

            # Red flags
            if ats.get('red_flags_found'):
                st.markdown("**Issues to Fix:**")
                for flag in ats['red_flags_found']:
                    st.markdown(f"{SYMBOLS['warning']} {flag}")

    else:
        st.info(f"{SYMBOLS['info']} Process your resume to see detailed analysis here.")

def display_export_options():
    """Display export and download options"""
    st.markdown(f"### {SYMBOLS['download']} Export Options")

    if 'enhanced_resume' in st.session_state and st.session_state.enhanced_resume:
        col1, col2 = st.columns(2)

        with col1:
            st.markdown("#### Enhanced Resume")

            # Text format download
            if st.button(f"{SYMBOLS['download']} Download as TXT", key="download_resume_txt"):
                download_as_txt('enhanced_resume')

            # Create DOCX download
            if st.button(f"{SYMBOLS['download']} Generate DOCX", key="generate_docx"):
                generate_docx_resume()

        with col2:
            st.markdown("#### Cover Letter")

            if 'cover_letter' in st.session_state and st.session_state.cover_letter:
                cover_letter = st.session_state.cover_letter

                # Select cover letter style
                style = st.selectbox(
                    "Select Style",
                    options=list(cover_letter.keys()),
                    format_func=lambda x: x.replace('_', ' ').title()
                )

                # Show selected cover letter
                if style in cover_letter:
                    st.text_area(
                        f"{style.replace('_', ' ').title()} Cover Letter",
                        cover_letter[style],
                        height=300,
                        disabled=True
                    )

                    if st.button(f"{SYMBOLS['download']} Download Cover Letter", key=f"download_cl_{style}"):
                        download_cover_letter(cover_letter[style], style)
            else:
                st.info(f"{SYMBOLS['info']} Complete the questionnaire to generate cover letters.")

    else:
        st.info(f"{SYMBOLS['info']} Process your resume to access export options.")

def download_as_txt(content_type):
    """Download content as text file"""
    if content_type == 'enhanced_resume' and 'enhanced_resume' in st.session_state:
        content = st.session_state.enhanced_resume.get('enhanced_full_content', '')
        filename = "enhanced_resume.txt"
    else:
        st.error("No content available for download")
        return

    # Create download link
    st.download_button(
        label=f"{SYMBOLS['download']} Download {filename}",
        data=content,
        file_name=filename,
        mime="text/plain"
    )

def download_cover_letter(content, style):
    """Download cover letter"""
    filename = f"cover_letter_{style}.txt"

    st.download_button(
        label=f"{SYMBOLS['download']} Download {filename}",
        data=content,
        file_name=filename,
        mime="text/plain"
    )

def generate_docx_resume():
    """Generate DOCX version of enhanced resume"""
    try:
        from docx import Document
        from docx.shared import Inches

        if 'enhanced_resume' not in st.session_state:
            st.error("No enhanced resume available")
            return

        enhanced = st.session_state.enhanced_resume

        # Create new document
        doc = Document()

        # Add title
        title = doc.add_heading('Enhanced Resume', 0)

        # Add professional summary
        if enhanced.get('enhanced_summary'):
            doc.add_heading('Professional Summary', level=1)
            doc.add_paragraph(enhanced['enhanced_summary'])

        # Add skills
        if enhanced.get('enhanced_skills'):
            doc.add_heading('Core Competencies', level=1)
            skills = enhanced['enhanced_skills']

            all_skills = []
            for category, skill_list in skills.items():
                if skill_list and category != 'suggested_additions':
                    all_skills.extend(skill_list)

            if all_skills:
                doc.add_paragraph(', '.join(all_skills))

        # Add experience
        if enhanced.get('enhanced_experience'):
            doc.add_heading('Professional Experience', level=1)
            for exp in enhanced['enhanced_experience']:
                # Add role and company
                doc.add_heading(f"{exp.get('role', 'Role')} - {exp.get('company', 'Company')}", level=2)

                # Add description points
                for desc in exp.get('enhanced_description', []):
                    p = doc.add_paragraph(desc, style='List Bullet')

        # Save to bytes
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)

        # Download button
        st.download_button(
            label=f"{SYMBOLS['download']} Download Enhanced Resume (DOCX)",
            data=doc_bytes.getvalue(),
            file_name="enhanced_resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.success(f"{SYMBOLS['check']} DOCX file generated successfully!")

    except ImportError:
        st.error(f"{SYMBOLS['warning']} python-docx library not available. Please install it for DOCX export.")
    except Exception as e:
        st.error(f"{SYMBOLS['warning']} Error generating DOCX: {str(e)}")
