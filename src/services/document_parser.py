"""
Document Parser Service
Uses only free libraries: spaCy, HuggingFace Transformers
"""

import spacy
from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
import docx
import PyPDF2
import pdfplumber
import io
import re
import logging
from typing import Dict, List, Any, Optional

class DocumentParser:
    """Parse documents and extract information using free NLP libraries"""

    def __init__(self):
        """Initialize the parser with free models"""
        self.setup_logging()
        self.load_models()

    def setup_logging(self):
        """Setup logging"""
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    def load_models(self):
        """Load free NLP models"""
        try:
            # Load spaCy model (free)
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                self.logger.warning("spaCy model not found, downloading...")
                spacy.cli.download("en_core_web_sm")
                self.nlp = spacy.load("en_core_web_sm")

            # Load HuggingFace NER model (free)
            model_name = "dbmdz/bert-large-cased-finetuned-conll03-english"
            self.ner_pipeline = pipeline(
                "ner", 
                model=model_name,
                aggregation_strategy="simple",
                device=-1  # Use CPU (free)
            )

            self.logger.info("✅ All models loaded successfully")

        except Exception as e:
            self.logger.error(f"❌ Error loading models: {e}")
            # Fallback to basic regex parsing
            self.nlp = None
            self.ner_pipeline = None

    def extract_content(self, uploaded_file) -> str:
        """Extract text content from uploaded file"""
        try:
            file_type = uploaded_file.type
            self.logger.info(f"Processing file type: {file_type}")

            # Reset file pointer
            uploaded_file.seek(0)

            if "application/pdf" in file_type:
                return self._extract_from_pdf(uploaded_file)
            elif "wordprocessingml" in file_type or "docx" in file_type:
                return self._extract_from_docx(uploaded_file)
            elif "text/plain" in file_type:
                return self._extract_from_txt(uploaded_file)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            self.logger.error(f"Content extraction failed: {e}")
            raise Exception(f"❌ Content extraction failed: {str(e)}")

    def _extract_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file using multiple methods"""
        text = ""

        try:
            # Method 1: pdfplumber (better for complex layouts)
            uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            if text.strip():
                return text

        except Exception as e:
            self.logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")

        try:
            # Method 2: PyPDF2 (fallback)
            uploaded_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        except Exception as e:
            self.logger.error(f"PyPDF2 also failed: {e}")
            raise Exception("Could not extract text from PDF")

        return text

    def _extract_from_docx(self, uploaded_file) -> str:
        """Extract text from DOCX file"""
        try:
            uploaded_file.seek(0)
            doc = docx.Document(uploaded_file)
            text = ""

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"

            return text

        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}")

    def _extract_from_txt(self, uploaded_file) -> str:
        """Extract text from TXT file"""
        try:
            uploaded_file.seek(0)
            content = uploaded_file.read()

            # Try to decode with different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue

            # If all fail, use errors='ignore'
            return content.decode('utf-8', errors='ignore')

        except Exception as e:
            raise Exception(f"TXT extraction failed: {str(e)}")

    def extract_entities(self, text: str) -> Dict[str, Any]:
        """Extract named entities and structured information"""
        entities = {
            'personal_info': {},
            'contact_info': {},
            'skills': [],
            'experience': [],
            'education': [],
            'projects': [],
            'organizations': [],
            'locations': []
        }

        try:
            # Use HuggingFace NER if available
            if self.ner_pipeline:
                ner_results = self.ner_pipeline(text)
                self._process_ner_results(ner_results, entities)

            # Use spaCy if available
            if self.nlp:
                doc = self.nlp(text)
                self._process_spacy_results(doc, entities)

            # Regex-based extraction (always runs as backup)
            self._extract_with_regex(text, entities)

        except Exception as e:
            self.logger.error(f"Entity extraction error: {e}")
            # Fall back to regex only
            self._extract_with_regex(text, entities)

        return entities

    def _process_ner_results(self, ner_results: List, entities: Dict):
        """Process HuggingFace NER results"""
        for entity in ner_results:
            entity_type = entity['entity_group']
            entity_text = entity['word']

            if entity_type == 'PER':
                if 'name' not in entities['personal_info']:
                    entities['personal_info']['name'] = entity_text
            elif entity_type == 'ORG':
                entities['organizations'].append(entity_text)
            elif entity_type == 'LOC':
                entities['locations'].append(entity_text)

    def _process_spacy_results(self, doc, entities: Dict):
        """Process spaCy NLP results"""
        for ent in doc.ents:
            if ent.label_ == 'PERSON' and 'name' not in entities['personal_info']:
                entities['personal_info']['name'] = ent.text
            elif ent.label_ == 'ORG':
                entities['organizations'].append(ent.text)
            elif ent.label_ in ['GPE', 'LOC']:
                entities['locations'].append(ent.text)

    def _extract_with_regex(self, text: str, entities: Dict):
        """Extract information using regex patterns"""

        # Email extraction
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text, re.IGNORECASE)
        if emails:
            entities['contact_info']['email'] = emails[0]

        # Phone extraction
        phone_patterns = [
            r'(\+?1?[-\s]?)?\(?([0-9]{3})\)?[-\s]?([0-9]{3})[-\s]?([0-9]{4})',
            r'(\+?\d{1,3}[-\s]?)?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{4}',
            r'\b\d{3}[-\.]\d{3}[-\.]\d{4}\b'
        ]

        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                entities['contact_info']['phone'] = ''.join(phones[0]) if isinstance(phones[0], tuple) else phones[0]
                break

        # Skills extraction (common technical skills)
        skill_keywords = [
            'Python', 'Java', 'JavaScript', 'React', 'Node.js', 'SQL', 'MongoDB',
            'Machine Learning', 'Data Science', 'AI', 'TensorFlow', 'PyTorch',
            'AWS', 'Docker', 'Kubernetes', 'Git', 'Linux', 'HTML', 'CSS',
            'Project Management', 'Leadership', 'Communication', 'Problem Solving',
            'Teamwork', 'Critical Thinking', 'Time Management'
        ]

        found_skills = []
        text_lower = text.lower()
        for skill in skill_keywords:
            if skill.lower() in text_lower:
                found_skills.append(skill)

        entities['skills'] = list(set(found_skills))  # Remove duplicates

        # Experience extraction (companies and roles)
        experience_patterns = [
            r'(Software Engineer|Data Scientist|Project Manager|Developer|Analyst|Manager|Director|Lead|Senior|Junior)\s+at\s+([A-Za-z0-9\s&]+)',
            r'([A-Za-z0-9\s&]+)\s*[-–]\s*(Software Engineer|Data Scientist|Project Manager|Developer|Analyst|Manager|Director|Lead|Senior|Junior)'
        ]

        for pattern in experience_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            for match in matches:
                if isinstance(match, tuple) and len(match) == 2:
                    role, company = match
                    entities['experience'].append({
                        'role': role.strip(),
                        'company': company.strip()
                    })

        # Education extraction
        education_keywords = [
            'Bachelor', 'Master', 'PhD', 'University', 'College', 'Institute',
            'Computer Science', 'Engineering', 'Business', 'Mathematics',
            'B.S.', 'M.S.', 'MBA', 'B.A.', 'M.A.'
        ]

        education_found = []
        for keyword in education_keywords:
            if keyword.lower() in text.lower():
                education_found.append(keyword)

        if education_found:
            entities['education'] = education_found

        return entities

    def validate_extraction(self, entities: Dict) -> Dict[str, Any]:
        """Validate and clean extracted entities"""
        validation_report = {
            'valid': True,
            'warnings': [],
            'suggestions': []
        }

        # Check for essential information
        if not entities['contact_info'].get('email'):
            validation_report['warnings'].append("No email address found")
            validation_report['suggestions'].append("Consider adding your email address")

        if not entities['contact_info'].get('phone'):
            validation_report['warnings'].append("No phone number found")
            validation_report['suggestions'].append("Consider adding your phone number")

        if not entities['skills']:
            validation_report['warnings'].append("No skills detected")
            validation_report['suggestions'].append("Consider adding a skills section")

        if not entities['experience']:
            validation_report['warnings'].append("No work experience detected")
            validation_report['suggestions'].append("Consider adding work experience details")

        return validation_report
